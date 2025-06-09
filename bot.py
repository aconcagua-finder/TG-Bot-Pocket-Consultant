#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import logging
import asyncio
import json
import pickle
from datetime import datetime, timedelta
from typing import Dict, Optional
import re

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from telegram.constants import ParseMode
import openai
import requests
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import io

# Загружаем переменные окружения
load_dotenv()

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=getattr(logging, os.getenv('LOG_LEVEL', 'INFO'))
)
logger = logging.getLogger(__name__)

# Конфигурация
TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
PERPLEXITY_API_KEY = os.getenv('PERPLEXITY_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')

# Инициализация OpenRouter (используем OpenAI SDK с изменённым базовым URL)
from openai import AsyncOpenAI
openrouter_client = AsyncOpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=OPENROUTER_API_KEY
)

# Хранилище для лимитов пользователей
user_limits = {}
DAILY_QUESTION_LIMIT = 10
DAILY_DOCUMENT_LIMIT = 10

class UserLimits:
    def __init__(self):
        self.questions_count = 0
        self.documents_count = 0
        self.last_reset = datetime.now().date()
    
    def reset_if_needed(self):
        today = datetime.now().date()
        if today > self.last_reset:
            self.questions_count = 0
            self.documents_count = 0
            self.last_reset = today
    
    def can_ask_question(self) -> bool:
        self.reset_if_needed()
        return self.questions_count < DAILY_QUESTION_LIMIT
    
    def can_process_document(self) -> bool:
        self.reset_if_needed()
        return self.documents_count < DAILY_DOCUMENT_LIMIT
    
    def increment_questions(self):
        self.reset_if_needed()
        self.questions_count += 1
    
    def increment_documents(self):
        self.reset_if_needed()
        self.documents_count += 1

def save_user_limits():
    """Сохраняет лимиты пользователей в файл"""
    try:
        with open("user_limits.pkl", "wb") as f:
            pickle.dump(user_limits, f)
    except Exception as e:
        logger.error(f"Error saving user limits: {e}")

def load_user_limits():
    """Загружает лимиты пользователей из файла"""
    global user_limits
    try:
        if os.path.exists("user_limits.pkl"):
            with open("user_limits.pkl", "rb") as f:
                user_limits = pickle.load(f)
            logger.info(f"Loaded user limits for {len(user_limits)} users")
        else:
            user_limits = {}
            logger.info("No existing user limits file found, starting fresh")
    except Exception as e:
        logger.error(f"Error loading user limits: {e}")
        user_limits = {}

def get_user_limits(user_id: int) -> UserLimits:
    if user_id not in user_limits:
        user_limits[user_id] = UserLimits()
    return user_limits[user_id]

def extract_text_from_file(file_content: bytes, file_extension: str) -> str:
    """Извлекает текст из файла в зависимости от его расширения"""
    try:
        if file_extension == '.txt':
            return file_content.decode('utf-8', errors='ignore')
        
        elif file_extension == '.pdf':
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_extension == '.docx':
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            return "Неподдерживаемый формат файла."
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_extension}: {e}")
        return f"Ошибка при извлечении текста из файла формата {file_extension}."

def log_user_action(user_id: int, username: str, action: str, query: str = ""):
    """Логирование действий пользователей"""
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "user_id": user_id,
        "username": username,
        "action": action,
        "query": query[:100] + "..." if len(query) > 100 else query
    }
    
    # Записываем в файл логов
    with open("user_logs.jsonl", "a", encoding="utf-8") as f:
        f.write(json.dumps(log_entry, ensure_ascii=False) + "\n")
    
    logger.info(f"User {user_id} ({username}): {action}")

def markdown_to_html(text: str) -> str:
    """Конвертирует markdown в HTML для Telegram"""
    # Удаляем или заменяем сложную разметку
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)  # **bold** -> <b>bold</b>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)      # *italic* -> <i>italic</i>
    text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)  # `code` -> <code>code</code>
    text = re.sub(r'###\s*(.*)', r'<b>\1</b>', text)     # ### Header -> <b>Header</b>
    text = re.sub(r'##\s*(.*)', r'<b>\1</b>', text)      # ## Header -> <b>Header</b>
    text = re.sub(r'#\s*(.*)', r'<b>\1</b>', text)       # # Header -> <b>Header</b>
    
    # Убираем только номера источников в квадратных скобках от [1] до [20]
    text = re.sub(r'\[([1-9]|1[0-9]|20)\]', '', text)
    
    # Убираем оставшуюся markdown разметку  
    text = re.sub(r'[\[\]()]', '', text)
    
    # Нормализуем переносы строк (сохраняем структуру)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Убираем лишние пустые строки
    text = re.sub(r'[ \t]+', ' ', text)  # Убираем только лишние пробелы и табы, НЕ переносы строк
    
    return text.strip()

def get_main_keyboard():
    """Создание основной клавиатуры"""
    keyboard = [
        [InlineKeyboardButton("❓ Задать вопрос", callback_data="ask_question")],
        [InlineKeyboardButton("📄 Анализ документа", callback_data="analyze_document")],
        [InlineKeyboardButton("✍️ Создать документ", callback_data="create_document")],
        [InlineKeyboardButton("ℹ️ Справка", callback_data="help")]
    ]
    return InlineKeyboardMarkup(keyboard)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /start"""
    user = update.effective_user
    
    welcome_message = (
        f"👋 Добро пожаловать, {user.first_name}!\n\n"
        
        "🤖 <b>Карманный Консультант</b> — ваш AI-помощник по юридическим вопросам РФ\n\n"
        
        "🎯 <b>Что я умею:</b>\n"
        "• ❓ <b>Отвечать на правовые вопросы</b> с использованием актуальной информации\n"
        "• 📄 <b>Анализировать документы</b> (.txt, .docx, .pdf)\n"
        "• ✍️ <b>Создавать юридические документы</b> по вашим требованиям\n\n"
        
        "📊 <b>Дневные лимиты:</b>\n"
        "• 10 вопросов\n"
        "• 10 операций с документами\n\n"
        
        "💡 Выберите нужное действие:"
    )
    
    await update.message.reply_text(
        welcome_message,
        parse_mode=ParseMode.HTML,
        reply_markup=get_main_keyboard()
    )

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик нажатий кнопок"""
    query = update.callback_query
    await query.answer()
    
    user = query.from_user
    
    if query.data == "ask_question":
        await query.edit_message_text(
            "❓ <b>Задайте ваш юридический вопрос</b>\n\n"
            "Опишите вашу ситуацию подробно. Чем больше деталей вы предоставите, "
            "тем более точный и полезный ответ вы получите.\n\n"
            "📝 <i>Просто напишите ваш вопрос в следующем сообщении</i>",
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("◀️ Назад", callback_data="back_to_main")]])
        )
        context.user_data['waiting_for'] = 'question'
    
    elif query.data == "analyze_document":
        await query.edit_message_text(
            "📄 <b>Анализ документа</b>\n\n"
            "📎 Отправьте документ для анализа\n\n"
            "✅ <b>Поддерживаемые форматы:</b>\n"
            "• .txt — текстовые файлы\n"
            "• .docx — документы Word\n"
            "• .pdf — PDF документы\n\n"
            "⚠️ <b>Ограничения:</b>\n"
            "• Максимальный размер: 20 МБ\n\n"
            "🔍 <b>Что получите:</b> Подробный анализ с выделением ключевых моментов, рисков и рекомендаций",
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("◀️ Назад", callback_data="back_to_main")]])
        )
        context.user_data['waiting_for'] = 'analyze_document'
    
    elif query.data == "create_document":
        await query.edit_message_text(
            "✍️ <b>Создание документа</b>\n\n"
            "📝 <b>Шаг 1:</b> Опишите какой документ вам нужен\n\n"
            "💡 <b>Примеры запросов:</b>\n"
            "• Договор аренды квартиры на 1 год\n"
            "• Претензия к поставщику за некачественный товар\n"
            "• Доверенность на получение документов\n"
            "• Заявление на отпуск\n"
            "• Уведомление о расторжении договора\n"
            "• Трудовой договор для программиста\n"
            "• Договор займа между физлицами\n\n"
            "📤 <b>Результат:</b> Готовый документ в формате .docx\n\n"
            "✏️ <i>Напишите подробно что за документ вам нужен</i>",
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("◀️ Назад", callback_data="back_to_main")]])
        )
        context.user_data['waiting_for'] = 'create_instructions'
    
    elif query.data == "help":
        await query.edit_message_text(
            "ℹ️ <b>Справка по использованию</b>\n\n"
            
            "🤖 <b>О боте:</b>\n"
            "Карманный Консультант — AI-помощник для решения юридических вопросов "
            "в рамках законодательства Российской Федерации.\n\n"
            
            "🎯 <b>Возможности:</b>\n\n"
            
            "❓ <b>Вопросы и ответы</b>\n"
            "• Правовые консультации\n"
            "• Разъяснения законов\n"
            "• Помощь в спорных ситуациях\n\n"
            
            "📄 <b>Анализ документов</b>\n"
            "• Изучение договоров\n"
            "• Проверка юридических документов\n"
            "• Выявление рисков и недочетов\n\n"
            
            "✍️ <b>Создание документов</b>\n"
            "• Договоры любых типов\n"
            "• Претензии и заявления\n"
            "• Доверенности и уведомления\n"
            "• Трудовые и гражданские документы\n\n"
            
            "⚠️ <b>Важно помнить:</b>\n"
            "Бот предоставляет информационную помощь. Для принятия серьезных "
            "юридических решений обязательно консультируйтесь с квалифицированным юристом.\n\n"
            
            "📞 <b>Поддержка:</b> @your_support",
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("◀️ Назад", callback_data="back_to_main")]])
        )
    
    elif query.data == "back_to_main":
        # Очищаем состояние пользователя
        context.user_data.clear()
        
        await query.edit_message_text(
            "🏠 <b>Главное меню</b>\n\n"
            "Выберите нужное действие:",
            parse_mode=ParseMode.HTML,
            reply_markup=get_main_keyboard()
        )

async def ask_perplexity(question: str) -> str:
    """Запрос к Perplexity API"""
    url = "https://api.perplexity.ai/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }
    
    # Добавляем контекст для юридических вопросов
    system_prompt = (
        "Ты — профессиональный юрист-консультант, специализирующийся на российском праве. "
        "Отвечай только на вопросы, связанные с правовой системой Российской Федерации. "
        "Если вопрос не касается российского права или юридических вопросов, "
        "вежливо откажись отвечать и предложи задать юридический вопрос. "
        "Давай подробные, профессиональные ответы со ссылками на нормативные акты где возможно."
    )
    
    data = {
        "model": "sonar-pro",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": question}
        ],
        "max_tokens": 2000,
        "temperature": 0.2
    }
    
    try:
        response = requests.post(url, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        
        result = response.json()
        return result['choices'][0]['message']['content']
    
    except Exception as e:
        logger.error(f"Perplexity API error: {e}")
        return "Извините, произошла ошибка при обработке вашего запроса. Попробуйте позже."

async def ask_chatgpt(prompt: str, document_content: str = "") -> str:
    """Запрос к OpenAI через OpenRouter API"""
    try:
        if document_content:
            full_prompt = f"{prompt}\n\nДокумент для обработки:\n{document_content}"
        else:
            full_prompt = prompt
        
        # Определяем системный промпт в зависимости от типа задачи
        if "документовед" in prompt or "Создай" in prompt:
            system_content = """Ты — опытный юрист-документовед, специализирующийся на создании юридических документов по российскому праву.

КРИТИЧЕСКИ ВАЖНО для создания документов:
- Создавай ПОЛНЫЕ, готовые к использованию документы
- Используй корректную юридическую терминологию РФ
- Включай все обязательные реквизиты и разделы
- Структурируй документы логично с нумерацией
- Добавляй поля для заполнения (ФИО, даты, адреса)
- Соблюдай действующее законодательство РФ
- Указывай ссылки на законы где уместно
- Добавляй места для подписей и печатей

КРИТИЧЕСКИ ВАЖНО ДЛЯ ФОРМАТИРОВАНИЯ:
- НЕ используй markdown-разметку (никаких #, **, _, и т.д.)
- НЕ используй символы # для заголовков
- НЕ используй ** для выделения жирным
- Просто пиши обычный текст без специальных символов
- Заголовки пиши заглавными буквами
- Для выделения используй только заглавные буквы или отступы

Работай только с юридическими документами РФ."""
        else:
            system_content = "Ты — профессиональный юрист, специализирующийся на анализе юридических документов по российскому праву. Работай только с юридическими документами. Если документ не связан с правовыми вопросами, вежливо откажись его обрабатывать и предложи прислать юридический документ."
        
        response = await openrouter_client.chat.completions.create(
            model="openai/gpt-4o-mini",  # Используем OpenAI модель через OpenRouter
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=3000,
            temperature=0.7,
            extra_headers={
                "HTTP-Referer": "https://pocket-consultant.ru",
                "X-Title": "Pocket Consultant Bot"
            }
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        logger.error(f"OpenRouter API error: {e}")
        return "Извините, произошла ошибка при обработке документа. Попробуйте позже."

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик текстовых сообщений"""
    user = update.effective_user
    message_text = update.message.text
    waiting_for = context.user_data.get('waiting_for')
    
    if waiting_for == 'create_instructions':
        # Сохраняем инструкции пользователя для создания документа
        context.user_data['create_instructions'] = message_text
        log_user_action(user.id, user.username or "Unknown", "create_instructions", message_text)
        
        # Сразу создаем документ
        user_limits_obj = get_user_limits(user.id)
        
        if not user_limits_obj.can_process_document():
            await update.message.reply_text(
                "⛔ Вы достигли дневного лимита операций с документами (10 в день).\n\n"
                "💡 <b>Что можно сделать:</b>\n"
                "• Обратиться завтра — лимит обнулится\n"
                "• Воспользоваться полной версией на "
                "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>",
                parse_mode=ParseMode.HTML,
                reply_markup=get_main_keyboard()
            )
            return

        await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
        
        # Отправляем сообщение-заглушку
        waiting_message = await update.message.reply_text(
            "⏳ Создаю документ по вашим требованиям...\n"
            "Это может занять до 30 секунд."
        )
        
        # Продолжаем показывать статус "печатает" во время обработки
        async def keep_typing():
            while True:
                await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
                await asyncio.sleep(4)  # Обновляем каждые 4 секунды
        
        # Запускаем индикатор печати в фоне
        typing_task = asyncio.create_task(keep_typing())
        
        try:
            # Формируем промпт для создания документа
            prompt = f"""Ты опытный юрист-документовед. Создай юридический документ согласно требованиям пользователя.

ТРЕБОВАНИЯ ПОЛЬЗОВАТЕЛЯ: {message_text}

ВАЖНЫЕ ТРЕБОВАНИЯ К СОЗДАНИЮ:
1. Создай ПОЛНЫЙ юридический документ со всеми необходимыми реквизитами
2. Используй корректную юридическую терминологию РФ
3. Включи все обязательные разделы для данного типа документа
4. Добавь необходимые поля для заполнения (ФИО, даты, адреса и т.д.)
5. Соблюдай действующее законодательство РФ
6. Структурируй документ логично с нумерацией пунктов
7. Добавь места для подписей и печатей где требуется
8. Укажи ссылки на соответствующие статьи законов где это уместно

КРИТИЧЕСКИ ВАЖНО ДЛЯ ФОРМАТИРОВАНИЯ:
- НЕ используй markdown-разметку (никаких #, **, _, и т.д.)
- НЕ используй символы # для заголовков
- НЕ используй ** для выделения жирным
- Просто пиши обычный текст без специальных символов
- Заголовки пиши заглавными буквами
- Для выделения используй только заглавные буквы или отступы

СТРУКТУРА ОТВЕТА:
- Заголовок документа (заглавными буквами)
- Все необходимые реквизиты
- Основной текст с пронумерованными пунктами
- Заключительная часть
- Места для подписей

Создай готовый к использованию документ БЕЗ markdown-разметки:"""
            
            logger.info(f"Creating document for user {user.id}: {message_text}")
            
            # Запрос к ChatGPT для создания документа
            response = await ask_chatgpt(prompt, "")
            logger.info(f"Document created successfully for user {user.id}")
            user_limits_obj.increment_documents()
            save_user_limits()
            
        finally:
            # Останавливаем индикатор печати
            typing_task.cancel()
            try:
                await typing_task
            except asyncio.CancelledError:
                pass
            
            # Удаляем сообщение-заглушку
            try:
                await waiting_message.delete()
            except Exception:
                pass  # Игнорируем ошибки удаления
        
        # Создаем и отправляем файл документа
        try:
            if response:
                # Создаем документ Word
                from docx import Document as DocxDocument
                doc = DocxDocument()
                
                # Разбиваем текст на параграфы и добавляем в документ
                paragraphs = response.split('\n\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        # Проверяем, является ли это заголовком
                        is_header = False
                        
                        # Заголовок, если весь текст заглавными буквами и короткий
                        if (paragraph_text.isupper() and len(paragraph_text) < 150) or \
                           any(word in paragraph_text.upper() for word in ['ДОГОВОР', 'СОГЛАШЕНИЕ', 'ЗАЯВЛЕНИЕ', 'ПРЕТЕНЗИЯ', 'ДОВЕРЕННОСТЬ', 'УВЕДОМЛЕНИЕ', 'ПРИКАЗ', 'РАСПОРЯЖЕНИЕ']):
                            is_header = True
                        
                        # Заголовок, если строка начинается с номера и содержит ключевые слова
                        elif any(paragraph_text.strip().upper().startswith(pattern) for pattern in [
                            '1.', '2.', '3.', '4.', '5.', 'I.', 'II.', 'III.', 'IV.', 'V.',
                            'ПРЕДМЕТ', 'ПРАВА И ОБЯЗАННОСТИ', 'ОТВЕТСТВЕННОСТЬ', 'СРОКИ', 'ЗАКЛЮЧИТЕЛЬНЫЕ',
                            'СТОРОНЫ ДОГОВОРА', 'ЦЕНА', 'ПОРЯДОК РАСЧЕТОВ', 'ФОРС-МАЖОР'
                        ]):
                            is_header = True
                        
                        if is_header:
                            p = doc.add_paragraph(paragraph_text.strip())
                            p.alignment = 1 if len(paragraph_text) < 100 else 0  # Центр для коротких заголовков
                            # Делаем заголовок жирным
                            for run in p.runs:
                                run.bold = True
                        else:
                            doc.add_paragraph(paragraph_text.strip())
                
                # Сохраняем в буфер
                import io
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_content = docx_buffer.getvalue()
                
                # Генерируем имя файла
                doc_type = "документ"
                if "договор" in message_text.lower():
                    doc_type = "договор"
                elif "заявление" in message_text.lower():
                    doc_type = "заявление"
                elif "претензия" in message_text.lower():
                    doc_type = "претензия"
                elif "доверенность" in message_text.lower():
                    doc_type = "доверенность"
                elif "уведомление" in message_text.lower():
                    doc_type = "уведомление"
                
                filename = f"{doc_type}.docx"
                
                # Отправляем документ
                await update.message.reply_document(
                    document=io.BytesIO(docx_content),
                    filename=filename,
                    caption=(
                        "✅ <b>Документ создан успешно!</b>\n\n"
                        "📄 Готовый юридический документ в формате Word\n\n"
                        "⚠️ <b>Важно:</b> Обязательно проверьте и адаптируйте документ под вашу конкретную ситуацию. "
                        "Перед использованием рекомендуется консультация с юристом.\n\n"
                        "🌟 Полная профессиональная версия доступна на "
                        "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>"
                    ),
                    parse_mode=ParseMode.HTML
                )
            else:
                raise Exception("Не удалось создать документ")
                
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}")
            await update.message.reply_text(
                "❌ <b>Ошибка при создании документа</b>\n\n"
                "Не удалось создать документ. Попробуйте:\n"
                "• Переформулировать запрос более подробно\n"
                "• Указать конкретный тип документа\n"
                "• Обратиться позже\n\n"
                "💡 Если проблема повторяется, свяжитесь с поддержкой.",
                parse_mode=ParseMode.HTML,
                reply_markup=get_main_keyboard()
            )
            return
        
        # Очищаем данные пользователя и возвращаем в главное меню
        context.user_data.clear()
        await update.message.reply_text(
            "Выберите следующее действие:",
            reply_markup=get_main_keyboard()
        )
        
    elif waiting_for == 'question':
        # Обработка вопроса через Perplexity
        user_limits_obj = get_user_limits(user.id)
        
        if not user_limits_obj.can_ask_question():
            await update.message.reply_text(
                "⛔ Вы достигли дневного лимита вопросов (10 в день).\n\n"
                "💡 <b>Что можно сделать:</b>\n"
                "• Обратиться завтра — лимит обнулится\n"
                "• Воспользоваться полной версией на "
                "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>",
                parse_mode=ParseMode.HTML,
                reply_markup=get_main_keyboard()
            )
            return
        
        log_user_action(user.id, user.username or "Unknown", "ask_question", message_text)
        
        # Показываем индикатор печати
        await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
        
        # Отправляем сообщение-заглушку
        waiting_message = await update.message.reply_text(
            "⏳ Подождите, пожалуйста, анализирую ваш запрос...\n"
            "Это может занять до 20 секунд."
        )
        
        # Продолжаем показывать статус "печатает" во время обработки
        async def keep_typing():
            while True:
                await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
                await asyncio.sleep(4)  # Обновляем каждые 4 секунды
        
        # Запускаем индикатор печати в фоне
        typing_task = asyncio.create_task(keep_typing())
        
        try:
            # Запрос к Perplexity
            answer = await ask_perplexity(message_text)
        finally:
            # Останавливаем индикатор печати
            typing_task.cancel()
            try:
                await typing_task
            except asyncio.CancelledError:
                pass
            
            # Удаляем сообщение-заглушку
            try:
                await waiting_message.delete()
            except Exception:
                pass  # Игнорируем ошибки удаления
        user_limits_obj.increment_questions()
        save_user_limits()  # Сохраняем изменения
        
        # Конвертируем markdown в HTML
        formatted_answer = markdown_to_html(answer)
        
        # Добавляем предупреждения
        final_answer = (
            f"{formatted_answer}\n\n"
            "⚠️ <b>Важно:</b> Это информационная консультация. "
            "Для принятия юридических решений обязательно проконсультируйтесь с квалифицированным юристом.\n\n"
            "🌟 Полная профессиональная версия доступна на "
            "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>"
        )
        
        # Проверяем длину сообщения (лимит Telegram ~4096 символов)
        if len(final_answer) > 4000:
            # Разбиваем на части
            parts = [final_answer[i:i+4000] for i in range(0, len(final_answer), 4000)]
            for part in parts:
                await update.message.reply_text(
                    part,
                    parse_mode=ParseMode.HTML,
                    disable_web_page_preview=True
                )
        else:
            await update.message.reply_text(
                final_answer,
                parse_mode=ParseMode.HTML,
                disable_web_page_preview=True
            )
        
        # Возвращаем главное меню
        await update.message.reply_text(
            "Выберите следующее действие:",
            reply_markup=get_main_keyboard()
        )
    
    else:
        # Если пользователь пишет без контекста
        await update.message.reply_text(
            "Для начала работы выберите нужное действие:",
            reply_markup=get_main_keyboard()
        )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик документов"""
    user = update.effective_user
    waiting_for = context.user_data.get('waiting_for')
    
    if waiting_for != 'analyze_document':
        await update.message.reply_text(
            "Для анализа документов сначала выберите соответствующее действие:",
            reply_markup=get_main_keyboard()
        )
        return
    
    user_limits_obj = get_user_limits(user.id)
    
    if not user_limits_obj.can_process_document():
        await update.message.reply_text(
            "⛔ Вы достигли дневного лимита обработки документов (10 в день).\n\n"
            "💡 <b>Что можно сделать:</b>\n"
            "• Обратиться завтра — лимит обнулится\n"
            "• Воспользоваться полной версией на "
            "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>",
            parse_mode=ParseMode.HTML,
            reply_markup=get_main_keyboard()
        )
        return
    
    document = update.message.document
    
    # Проверяем размер файла (20 МБ)
    if document.file_size > 20 * 1024 * 1024:
        await update.message.reply_text(
            "❌ Файл слишком большой. Максимальный размер: 20 МБ",
            reply_markup=get_main_keyboard()
        )
        return
    
    # Проверяем формат файла
    allowed_extensions = ['.txt', '.docx', '.pdf']
    file_extension = os.path.splitext(document.file_name)[1].lower()
    
    if file_extension not in allowed_extensions:
        await update.message.reply_text(
            "❌ Неподдерживаемый формат файла. Поддерживаются: .txt, .docx, .pdf",
            reply_markup=get_main_keyboard()
        )
        return
    
    log_user_action(user.id, user.username or "Unknown", "analyze_document_file", document.file_name)
    
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    
    # Отправляем сообщение-заглушку для документа
    waiting_message = await update.message.reply_text(
        "⏳ Подождите, пожалуйста, анализирую документ...\n"
        "Это может занять до 30 секунд."
    )
    
    # Продолжаем показывать статус "печатает" во время обработки
    async def keep_typing():
        while True:
            await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
            await asyncio.sleep(4)  # Обновляем каждые 4 секунды
    
    # Запускаем индикатор печати в фоне
    typing_task = asyncio.create_task(keep_typing())
    
    response = None  # Инициализируем переменную
    
    try:
        # Получаем файл
        file = await context.bot.get_file(document.file_id)
        file_content = await file.download_as_bytearray()
        
        # Извлекаем текст из файла
        document_text = extract_text_from_file(file_content, file_extension)
        
        # Проверяем что текст извлечен
        if not document_text or len(document_text.strip()) < 10:
            raise Exception("Не удалось извлечь текст из документа или документ слишком короткий")
        
        # Формируем запрос для ChatGPT для анализа
        prompt = "Проанализируй этот юридический документ. Выдели основные пункты, возможные риски и рекомендации."
        
        logger.info(f"Processing document analysis for user {user.id}, document length: {len(document_text)} chars")
        logger.info(f"Original document preview: {document_text[:200]}...")
        
        # Запрос к ChatGPT
        response = await ask_chatgpt(prompt, document_text[:8000])  # Ограничиваем длину
        logger.info(f"AI response preview: {response[:200]}..." if response else "No response received")
        user_limits_obj.increment_documents()
        save_user_limits()  # Сохраняем изменения
        
    finally:
        # Останавливаем индикатор печати
        typing_task.cancel()
        try:
            await typing_task
        except asyncio.CancelledError:
            pass
        
        # Удаляем сообщение-заглушку
        try:
            await waiting_message.delete()
        except Exception:
            pass  # Игнорируем ошибки удаления
    
    # Форматируем ответ
    try:
        if response is None:
            raise Exception("No response received")
        
        # Конвертируем markdown в HTML
        formatted_response = markdown_to_html(response)
        
        # Добавляем предупреждения
        final_response = (
            f"📄 <b>Анализ документа:</b>\n\n"
            f"{formatted_response}\n\n"
            "⚠️ <b>Важно:</b> Это автоматический анализ. "
            "Для принятия юридических решений обязательно проконсультируйтесь с квалифицированным юристом.\n\n"
            "🌟 Полная профессиональная версия доступна на "
            "<a href='https://pocket-consultant.ru'>pocket-consultant.ru</a>"
        )
        
        # Проверяем длину сообщения
        if len(final_response) > 4000:
            # Разбиваем на части
            parts = [final_response[i:i+4000] for i in range(0, len(final_response), 4000)]
            for part in parts:
                await update.message.reply_text(
                    part,
                    parse_mode=ParseMode.HTML,
                    disable_web_page_preview=True
                )
        else:
            await update.message.reply_text(
                final_response,
                parse_mode=ParseMode.HTML,
                disable_web_page_preview=True
            )
    
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        await update.message.reply_text(
            "❌ <b>Ошибка при обработке документа</b>\n\n"
            "Не удалось проанализировать документ. Возможные причины:\n"
            "• Документ поврежден или не содержит текст\n"
            "• Временная техническая проблема\n\n"
            "💡 Попробуйте повторить попытку позже.",
            parse_mode=ParseMode.HTML,
            reply_markup=get_main_keyboard()
        )
        return
    
    # Очищаем состояние и возвращаем главное меню
    context.user_data.clear()
    await update.message.reply_text(
        "Выберите следующее действие:",
        reply_markup=get_main_keyboard()
    )

async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик ошибок"""
    logger.error(f"Exception while handling an update: {context.error}")

def main():
    """Основная функция запуска бота"""
    if not TELEGRAM_BOT_TOKEN:
        logger.error("TELEGRAM_BOT_TOKEN not found in environment variables")
        return
    
    # Загружаем лимиты пользователей
    load_user_limits()
    
    # Создаем приложение
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    # Добавляем обработчики
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CallbackQueryHandler(button_handler))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    
    # Добавляем обработчик ошибок
    application.add_error_handler(error_handler)
    
    # Запускаем бота
    logger.info("Starting bot...")
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main() 